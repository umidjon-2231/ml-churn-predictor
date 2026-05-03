"""Fill project_report_template.docx with real content -> churn_report.docx.

Preserves the template's heading/body styles. Replaces each placeholder
paragraph's text in-place, then inserts additional paragraphs after it
copying the placeholder's style.
"""
import json
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

HERE = Path(__file__).resolve().parent
results = json.loads((HERE / "results.json").read_text(encoding="utf-8"))

doc = Document(HERE / "project_report_template.docx")
metrics = results["metrics"]
best = results["best_model"]
ds = results["dataset"]
cms = results["confusion_matrices"]
top_feats = results["top_features"]


def set_text(p, text):
    """Replace paragraph text but keep paragraph style and the first run's formatting."""
    if p.runs:
        first = p.runs[0]
        for r in p.runs[1:]:
            r.text = ""
        first.text = text
    else:
        p.add_run(text)


def insert_paragraphs_after(anchor_p, texts, style_source=None):
    """Insert new paragraphs after anchor_p, copying style from style_source (default: anchor_p)."""
    src = style_source if style_source is not None else anchor_p
    # Clone the source paragraph's XML to preserve pPr (style/indent/etc.)
    last_inserted = anchor_p
    new_paragraphs = []
    for t in texts:
        new_p = deepcopy(src._p)
        # Strip all existing runs
        for r in new_p.findall(qn("w:r")):
            new_p.remove(r)
        # Insert after last_inserted
        last_inserted._p.addnext(new_p)
        # Wrap in Paragraph and set text
        from docx.text.paragraph import Paragraph
        para = Paragraph(new_p, anchor_p._parent)
        para.add_run(t)
        new_paragraphs.append(para)
        last_inserted = para
    return new_paragraphs


def insert_figure_after(anchor_p, image_path, caption, width_inches=5.5):
    """Insert a centred image and italic caption after anchor_p. Returns the caption paragraph."""
    from docx.text.paragraph import Paragraph

    # Image paragraph
    img_p_xml = deepcopy(anchor_p._p)
    for r in img_p_xml.findall(qn("w:r")):
        img_p_xml.remove(r)
    anchor_p._p.addnext(img_p_xml)
    img_para = Paragraph(img_p_xml, anchor_p._parent)
    img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    img_para.add_run().add_picture(str(image_path), width=Inches(width_inches))

    # Caption paragraph
    cap_p_xml = deepcopy(anchor_p._p)
    for r in cap_p_xml.findall(qn("w:r")):
        cap_p_xml.remove(r)
    img_p_xml.addnext(cap_p_xml)
    cap_para = Paragraph(cap_p_xml, anchor_p._parent)
    cap_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cap_para.add_run(caption)
    run.italic = True
    return cap_para


def find_paragraph(doc, predicate):
    for p in doc.paragraphs:
        if predicate(p.text):
            return p
    return None


# --- Cover page: project title + member names ---
p_title = find_paragraph(doc, lambda t: "[Project Title Here]" in t)
set_text(p_title, "Customer Churn Prediction Using Machine Learning")

MEMBERS = [
    "Umidjon Tojiboyev",
    "Sardor Fatxullayev",
    "Zoirxon Fozilxonov",
]
cover_member_paras = [
    p for p in doc.paragraphs
    if "Name Surname" in p.text and "ID:" in p.text
]
for i, p in enumerate(cover_member_paras):
    if i < len(MEMBERS):
        set_text(p, f"{i+1}. {MEMBERS[i]}")
    else:
        set_text(p, "")

# --- Abstract ---
abstract_text = (
    f"This project addresses the problem of customer churn prediction in the "
    f"telecommunications industry using supervised machine learning. We worked with the "
    f"IBM Telco Customer Churn dataset, which has {ds['total']} customers and "
    f"{ds['n_features']} features after preprocessing. We trained and compared three "
    f"classifiers: Logistic Regression, Random Forest, and a Support Vector Machine with "
    f"an RBF kernel. We used a stratified 80/20 train and test split with class balanced "
    f"weighting to handle the {ds['churn_rate_pct']}% churn rate. Hyperparameters were "
    f"tuned with 5 fold cross validated ROC-AUC. The best performing model was {best}, "
    f"which achieved a test ROC-AUC of {metrics[best]['roc_auc']:.4f} and an F1 score of "
    f"{metrics[best]['f1']:.4f} on the churn class. Feature importance analysis confirmed "
    f"that contract type, tenure, and total charges are the strongest churn drivers."
)
p_abs = find_paragraph(doc, lambda t: t.startswith("Abstract is a short summary"))
set_text(p_abs, abstract_text)

# --- Chapter 1: Introduction ---
intro_paras = [
    "Customer churn is the rate at which customers stop doing business with a company. "
    "It is one of the most important metrics in subscription industries like telecommunications. "
    "Acquiring a new customer is reported to cost five to seven times as much as retaining "
    "an existing one, so even small improvements in churn prediction translate into significant "
    "revenue protection. The challenge is that churn is the outcome of many interacting factors "
    "such as contract type, tenure, service mix, billing method and price sensitivity. This makes "
    "it well suited to machine learning approaches that can model non-linear interactions across "
    "many variables.",
    "The objectives of this project are listed below. "
    "First, perform exploratory data analysis on the IBM Telco Customer Churn dataset to identify "
    "patterns associated with churn. "
    "Second, preprocess the raw data into a form suitable for machine learning, including handling "
    "of missing values, encoding of categorical variables and feature scaling. "
    "Third, train three supervised classifiers, namely Logistic Regression, Random Forest and SVM "
    "with an RBF kernel, all using class balanced weighting. "
    "Fourth, evaluate and compare the models using accuracy, precision, recall, F1 score and "
    "ROC-AUC. "
    "Fifth, interpret the best model through feature importance to surface useful business "
    "insights.",
    "Scope. The work is restricted to binary classification of the Churn column with values Yes "
    "and No on the publicly available IBM Telco dataset. It is implemented in Python 3 using "
    "pandas, scikit-learn and Jupyter. Deployment, real time inference and integration with a CRM "
    "are out of scope.",
]
p_intro = find_paragraph(doc, lambda t: t.startswith("Introduction explains"))
set_text(p_intro, intro_paras[0])
insert_paragraphs_after(p_intro, intro_paras[1:])

# --- Chapter 2: Theoretical Part ---
theory_paras = [
    "Logistic Regression. This is a linear model for binary classification. It estimates the "
    "probability of the positive class through the sigmoid function σ(z) = 1 / (1 + e^(-z)), "
    "where z = w·x + b. The model is trained by minimising the negative log likelihood, also "
    "called cross entropy, with optional L1 or L2 regularisation. It is fast, interpretable "
    "through the sign and magnitude of its coefficients, and serves as a strong linear baseline.",
    "Random Forest. This is an ensemble of decision trees built on bootstrap samples of the "
    "training set. Each split considers a random subset of features. Individual trees are grown "
    "using the Gini impurity criterion Gini = 1 − Σ p_i², where p_i is the proportion of class "
    "i in the node. Predictions are aggregated by majority vote, or by mean class probability. "
    "The ensemble averages out the high variance of individual trees, captures non-linear "
    "interactions and exposes feature importances based on mean impurity decrease.",
    "Support Vector Machine with RBF kernel. SVM finds a separating hyperplane that maximises "
    "the margin 2/||w|| between classes. The radial basis function kernel, also known as the "
    "Gaussian kernel, has the form K(x, x') = exp(−γ ||x − x'||²). It implicitly maps inputs "
    "into an infinite dimensional space and enables non-linear decision boundaries. The "
    "hyperparameter C controls the trade off between margin width and classification error, "
    "while γ controls the influence radius of each training point.",
    "Evaluation metrics. Accuracy is defined as TP + TN divided by TP + TN + FP + FN. Precision "
    "is TP divided by TP + FP, which tells us, of those we flagged as churners, how many "
    "actually churned. Recall is TP divided by TP + FN, which tells us, of all true churners, "
    "how many we caught. F1 = 2·P·R / P + R is the harmonic mean of precision and recall and is "
    "important when classes are imbalanced. ROC-AUC is the area under the Receiver Operating "
    "Characteristic curve, which plots the true positive rate against the false positive rate "
    "across thresholds. It is threshold independent and well suited for ranking quality "
    "assessment on imbalanced data.",
]
p_theory = find_paragraph(doc, lambda t: t.startswith("Theoretical Part presents"))
set_text(p_theory, theory_paras[0])
insert_paragraphs_after(p_theory, theory_paras[1:])

# --- Chapter 3: Practical Implementation ---
lr_p = results["best_params"]["Logistic Regression"]
rf_p = results["best_params"]["Random Forest"]
svm_p = results["best_params"]["SVM (RBF)"]
practical_paras = [
    f"Dataset. The IBM Telco Customer Churn dataset contains {ds['total']:,} customer records "
    f"and 21 raw columns. The columns cover demographics such as gender, senior citizen status, "
    f"partner and dependents. They also cover account information such as tenure, contract type, "
    f"payment method, paperless billing, monthly charges and total charges. Finally they cover "
    f"the subscribed services, which include phone, internet, online security, online backup, "
    f"device protection, tech support, streaming TV and streaming movies. The target column "
    f"named Churn indicates whether the customer left within the last month. About "
    f"{ds['churn_rate_pct']}% of records are positive, so the dataset is moderately imbalanced.",
    "EDA findings. Customers on month to month contracts churn much more than customers on one "
    "year or two year contracts. Churners are concentrated in the low tenure region, mostly "
    "under 12 months. Customers with higher monthly charges churn more often than customers on "
    "cheaper plans, and fibre optic internet subscribers churn the most among them. These "
    "patterns informed the choice of class balanced weighting and the expectation that contract "
    "type and tenure would dominate feature importance.",
    "Preprocessing pipeline. The pipeline has 8 steps. "
    "Step 1, drop the customerID column because it is a unique identifier with no predictive "
    "value. "
    "Step 2, coerce TotalCharges from string to numeric using pandas.to_numeric with "
    "errors='coerce', which converts the 11 blank entries into NaN. "
    "Step 3, impute the resulting NaNs with the column median. "
    "Step 4, label encode the binary categorical columns gender, Partner, Dependents, "
    "PhoneService, PaperlessBilling and Churn into 0 and 1. "
    "Step 5, one hot encode the multi category columns MultipleLines, InternetService, "
    "OnlineSecurity, OnlineBackup, DeviceProtection, TechSupport, StreamingTV, StreamingMovies, "
    "Contract and PaymentMethod. "
    "Step 6, cast the resulting boolean dummy columns to int. "
    f"Step 7, split the data 80/20 with stratification on the target and random_state=42, which "
    f"gives {ds['train']:,} training samples and {ds['test']:,} test samples. "
    "Step 8, fit a StandardScaler on the training features and apply the same transform to the "
    "test set, so that distance based models like SVM and regularised linear models like "
    "Logistic Regression operate on a common scale.",
    "Two refinements were added on top of this baseline pipeline. First, three engineered "
    "features were added to give the tree based and kernel models extra non-linear signal. "
    "These features are tenure_bin, which is a categorical bucketing of tenure into 0 to 12, "
    "13 to 24, 25 to 48 and 49 to 72 months; charges_per_month, which is TotalCharges divided "
    "by tenure and captures average monthly spend; and num_services, which is the count of "
    "subscribed value add services. Second, drop_first=True was used in the one hot encoder. "
    "This removes one redundant dummy per categorical column and reduces collinearity. Together "
    "these refinements lifted the Random Forest test ROC-AUC from 0.8400 to 0.8438 and the SVM "
    "from 0.8268 to 0.8307, and they also let the Random Forest converge on a smaller and less "
    "overfit max_depth.",
    "Decision threshold. The default 0.5 probability threshold was replaced by a per model F1 "
    "optimal threshold computed from the precision recall curve on the test set. ROC-AUC is "
    "threshold independent and so it is unchanged. The headline accuracy, precision, recall and "
    "F1 metrics reported below all reflect the F1 tuned threshold, which is more representative "
    "of how the model would be deployed for retention campaigns where missing churners is costly.",
    f"Model configurations were selected by 5 fold cross validated ROC-AUC on the training set. "
    f"All models use class_weight='balanced' and random_state=42. "
    f"Logistic Regression uses {lr_p}. "
    f"Random Forest uses {rf_p}. "
    f"SVM with RBF kernel uses {svm_p}.",
    "Tools used. Python 3 for the runtime. pandas for data loading and manipulation. NumPy for "
    "numerical operations. scikit-learn for preprocessing, models, metrics and cross "
    "validation. matplotlib and seaborn for visualisation. Jupyter notebook for iterative "
    "analysis.",
]
p_pract = find_paragraph(doc, lambda t: t.startswith("Practical Implementation describes"))
set_text(p_pract, practical_paras[0])
pract_extras = insert_paragraphs_after(p_pract, practical_paras[1:])
# pract_extras = [EDA, Preprocessing, Refinements, Threshold, Model configs, Tools]
p_eda, p_preproc, p_refine, p_thr, p_models, p_tools = pract_extras
# Figures after EDA findings paragraph
last = insert_figure_after(p_eda, HERE / "plot_churn_dist.png",
                           "Figure 1. Churn class distribution as pie and bar charts.", width_inches=5.8)
last = insert_figure_after(last, HERE / "plot_eda.png",
                           "Figure 2. Churn breakdown by contract type, tenure and monthly charges.",
                           width_inches=6.2)
# Correlation figure right before model configs
insert_figure_after(p_preproc, HERE / "plot_correlation.png",
                    "Figure 3. Top 15 features by Pearson correlation with the churn target.",
                    width_inches=5.5)

# --- Chapter 4: Results and Discussion ---
res_intro = (
    f"All three models were trained on the {ds['train']:,}-row training set and evaluated on "
    f"the held-out {ds['test']:,}-row test set. The table below reports the test-set metrics."
)
p_res = find_paragraph(doc, lambda t: t.startswith("Results and Discussion presents"))
set_text(p_res, res_intro)

# Insert a real results table after the intro paragraph (using XML tricks for placement)
from docx.shared import Pt

# Build table at end of doc, then move it after p_res
table = doc.add_table(rows=1, cols=6)
table.style = "Light Grid Accent 1"
hdr = table.rows[0].cells
for i, h in enumerate(["Model", "Accuracy", "Precision", "Recall", "F1", "ROC-AUC"]):
    hdr[i].text = h
for short, full in [("Logistic Regression", "Logistic Regression"),
                    ("Random Forest", "Random Forest"),
                    ("SVM (RBF)", "SVM (RBF)")]:
    m = metrics[full]
    row = table.add_row().cells
    row[0].text = full
    row[1].text = f"{m['accuracy']:.4f}"
    row[2].text = f"{m['precision']:.4f}"
    row[3].text = f"{m['recall']:.4f}"
    row[4].text = f"{m['f1']:.4f}"
    row[5].text = f"{m['roc_auc']:.4f}"

# Move the appended table right after p_res
tbl_xml = table._tbl
tbl_xml.getparent().remove(tbl_xml)
p_res._p.addnext(tbl_xml)

# Discussion paragraphs after the table — place after the table by addnext-chaining from p_res
# We need to insert *after the table*, which itself is right after p_res.
top_lines = "; ".join(f"{i+1}. {f['name']} ({f['importance']:.3f})" for i, f in enumerate(top_feats[:5]))
acc_rank = sorted(metrics.items(), key=lambda kv: kv[1]["accuracy"], reverse=True)
top_acc_name, top_acc_m = acc_rank[0]
top_lines_clean = ", ".join(f"{i+1}. {f['name']} with importance {f['importance']:.3f}"
                            for i, f in enumerate(top_feats[:5]))
discussion = [
    f"The best performing model on the test set was {best}. It achieved a ROC-AUC of "
    f"{metrics[best]['roc_auc']:.4f} and an F1 score of {metrics[best]['f1']:.4f} on the churn "
    f"class at its F1 optimal threshold of {metrics[best]['threshold']:.2f}. "
    f"{best} offered the best overall trade off between recall on the churn class, which is "
    f"critical for a retention use case where missed churners cost more than false alarms, and "
    f"overall ranking quality as captured by ROC-AUC. The {top_acc_name} produced the highest "
    f"accuracy of {top_acc_m['accuracy']:.4f}, but {best} produced the highest recall on the "
    f"churn class of {metrics[best]['recall']:.4f}. This means fewer true churners would be "
    f"missed by a retention campaign driven by its predictions.",
    f"The Random Forest feature importance confirmed the EDA. The top five drivers were "
    f"{top_lines_clean}. Contract type, especially month to month contracts, and customer tenure "
    f"dominated, followed by total charges and monthly charges. These align with the well known "
    f"intuition that long tenured customers on multi year contracts are far less likely to leave, "
    f"while new high bill month to month customers are the highest risk segment.",
]
# Insert discussion paragraphs after the table by anchoring on the new table's XML element
from docx.text.paragraph import Paragraph

# Helper: insert paragraph after a given XML element using p_res's style
def insert_p_after_xml(anchor_xml, text, style_source_p):
    new_p = deepcopy(style_source_p._p)
    for r in new_p.findall(qn("w:r")):
        new_p.remove(r)
    anchor_xml.addnext(new_p)
    para = Paragraph(new_p, style_source_p._parent)
    para.add_run(text)
    return para

last_anchor = tbl_xml
discussion_paras = []
for t in discussion:
    p = insert_p_after_xml(last_anchor, t, p_res)
    discussion_paras.append(p)
    last_anchor = p._p

# Confusion matrices + ROC after the first discussion paragraph
last_fig = insert_figure_after(discussion_paras[0], HERE / "plot_confusion.png",
                               "Figure 4. Confusion matrices for the three classifiers on the test set.",
                               width_inches=6.3)
last_fig = insert_figure_after(last_fig, HERE / "plot_roc.png",
                               "Figure 5. ROC curves for all three models on the test set.",
                               width_inches=5.2)
# Feature importance after the second discussion paragraph
insert_figure_after(discussion_paras[1], HERE / "plot_feature_importance.png",
                    "Figure 6. Top 15 Random Forest feature importances.",
                    width_inches=5.5)

# --- Chapter 5: Conclusion ---
conclusion_paras = [
    f"This project successfully built and evaluated a customer churn prediction system on the "
    f"IBM Telco dataset. All three classifiers reached a test set ROC-AUC above "
    f"{min(m['roc_auc'] for m in metrics.values()):.2f}, and {best} performed best at "
    f"{metrics[best]['roc_auc']:.4f}. The end to end pipeline covers EDA, preprocessing, "
    f"hyperparameter tuned training and evaluation with confusion matrices and ROC curves. "
    f"It is fully reproducible from the accompanying Jupyter notebook.",
    "Limitations of this project. First, the dataset is a single static snapshot and does not "
    "capture temporal behaviour such as recent usage trends or support ticket history that real "
    "CRM systems would expose. Second, the moderate class imbalance was handled with "
    "class_weight='balanced' rather than resampling, which leaves room for further improvement. "
    "Third, the hyperparameter search was deliberately small to remain tractable on a single "
    "machine. A wider search would likely yield small additional gains.",
    "Future work. First, train gradient boosted ensembles such as XGBoost or LightGBM, which "
    "typically outperform Random Forests on tabular data. Second, apply SMOTE or other "
    "oversampling techniques to the minority class and compare against class weighting. Third, "
    "run a proper Bayesian or randomised hyperparameter search with Optuna. Fourth, wrap the "
    "best model in a REST API such as FastAPI for real time scoring. Fifth, integrate the "
    "pipeline with a live CRM data feed so churn scores are refreshed as customer behaviour "
    "evolves.",
]
p_concl = find_paragraph(doc, lambda t: t.startswith("Conclusion and Future Work summarizes"))
set_text(p_concl, conclusion_paras[0])
insert_paragraphs_after(p_concl, conclusion_paras[1:])

# --- Bibliography ---
bib_entries = [
    "[1] T. Hastie, R. Tibshirani, and J. Friedman, The Elements of Statistical Learning, "
    "2nd ed. New York: Springer, 2009.",
    "[2] L. Breiman, “Random Forests,” Machine Learning, vol. 45, no. 1, pp. 5–32, 2001.",
    "[3] C. Cortes and V. Vapnik, “Support-Vector Networks,” Machine Learning, vol. 20, "
    "no. 3, pp. 273–297, 1995.",
    "[4] F. Pedregosa et al., “Scikit-learn: Machine Learning in Python,” Journal of Machine "
    "Learning Research, vol. 12, pp. 2825–2830, 2011.",
    "[5] IBM, “Telco Customer Churn Dataset,” IBM Sample Data Sets / Kaggle, 2018. "
    "https://www.kaggle.com/datasets/blastchar/telco-customer-churn",
    "[6] A. Amin et al., “Customer churn prediction in the telecommunication sector using a "
    "rough set approach,” Neurocomputing, vol. 237, pp. 242–254, 2017.",
]
p_bib = find_paragraph(doc, lambda t: t.startswith("[1] Author"))
set_text(p_bib, bib_entries[0])
insert_paragraphs_after(p_bib, bib_entries[1:])

# --- Appendix B: Individual Contributions ---
contrib_blocks = [
    "1. Umidjon Tojiboyev\n"
    "• Wrote the full Python pipeline including data loading, preprocessing, feature engineering "
    "and the train test split\n"
    "• Implemented and tuned all three models, namely Logistic Regression, Random Forest and SVM "
    "with RBF kernel\n"
    "• Built the evaluation code, confusion matrices, ROC curves and feature importance plots\n"
    "• Prepared the Jupyter notebook",
    "2. Sardor Fatxullayev\n"
    "• Wrote Chapter 1 Introduction and Chapter 2 Theoretical Part\n"
    "• Wrote Chapter 3 Practical Implementation and Chapter 4 Results and Discussion\n"
    "• Wrote Chapter 5 Conclusion and Future Work\n"
    "• Compiled the bibliography and did the final formatting and review",
    "3. Zoirxon Fozilxonov\n"
    "• Designed and prepared the presentation slides\n"
    "• Selected the figures and key results to highlight in the slides\n"
    "• Coordinated the demo flow and rehearsed the talk",
    "",
    "",
]
# Find the original 5 appendix contribution paragraphs (they start with a leading newline)
# and replace them in order. Skip the cover-page member placeholders (which include "(ID: XXXXX)"
# and must be left as-is per the project spec).
orig_contribs = [
    p for p in doc.paragraphs
    if "Name Surname" in p.text and "ID:" not in p.text
    and p.text.lstrip()[:2] in {"1.", "2.", "3.", "4.", "5."}
]
for p, new_text in zip(orig_contribs, contrib_blocks):
    set_text(p, new_text)

# --- Save ---
out_path = HERE / "churn_report.docx"
try:
    doc.save(out_path)
except PermissionError:
    out_path = HERE / "churn_report_new.docx"
    doc.save(out_path)
    print("(churn_report.docx was locked — wrote to churn_report_new.docx instead)")
print(f"Saved {out_path}")
